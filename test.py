# -*- coding: utf-8 -*-
# @File  : test.py
# @Author: Panbo
# @Date  : 2019/2/20
# @Desc  :
import docx
document = docx.Document('test.docx')


# 总共有多少段落
print('总共有 {} 个段落'.format(len(document.paragraphs)))
for paragraph in document.paragraphs:  # 遍历paragraph对象的列表
    # paragraph 为每个段落的对象 docx.text.paragraph.Paragraph object

    text = paragraph.text  # 每一段的文字内容
    print(text)
